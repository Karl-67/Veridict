"""
Routes for contract clause listing, editing, finding accept/dismiss,
and document annotations (anchored comments / suggestions).
"""
from __future__ import annotations

import html
import logging
import re
import uuid
from datetime import datetime
from pathlib import Path
from typing import Optional

import io

logger = logging.getLogger(__name__)

from fastapi import APIRouter, Body, Depends, HTTPException, status
from fastapi.responses import PlainTextResponse, StreamingResponse
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.orm import selectinload
from sqlalchemy.orm.attributes import flag_modified

from app.backend.db.models import (
    DocumentAnnotationRecord,
    FindingRecord,
    HumanReviewRecord,
    ParsedClauseRecord,
    RunRecord,
)
from app.backend.db.session import DbSession
from app.backend.models.schemas import DocumentAnnotationCreate, DocumentAnnotationOut
from app.backend.services.auth_service import require_auth

router = APIRouter(prefix="/api", tags=["editor"])


# ---------------------------------------------------------------------------
# Schemas
# ---------------------------------------------------------------------------

class CommentAnchorOut(BaseModel):
    annotation_id: str
    from_pos: int
    to_pos: int


class PendingSuggestionOut(BaseModel):
    finding_id: str
    severity: str
    description: str
    replacement_text: str


class DraftBlockOut(BaseModel):
    block_id: str
    clause_uid: str
    page_number: int
    style: str  # heading1 | heading2 | heading3 | body | list_item
    text: str           # current text (after accepted edits)
    original_text: str  # original parsed text
    marks: list[dict]   # [{type, from_pos, to_pos}]
    pending_suggestion: Optional[PendingSuggestionOut]
    comment_anchors: list[CommentAnchorOut]


class DocumentDraftOut(BaseModel):
    blocks: list[DraftBlockOut]
    revision: int


class ClauseOut(BaseModel):
    clause_uid: str
    page_number: int
    normalized_text: str
    bbox: list[float]


class ClauseEditBody(BaseModel):
    text: str
    plain_text: Optional[str] = None
    rich_text: Optional[list[dict]] = None
    page: Optional[int] = None
    rects: Optional[list[dict]] = None
    anchor_text: Optional[str] = None


class ContractEditsOut(BaseModel):
    edits: dict  # clause_uid → {text, edited_at}


class AcceptFindingBody(BaseModel):
    custom_text: Optional[str] = None  # override replacement text


class PdfRectOut(BaseModel):
    page: int
    x0: float
    top: float
    x1: float
    bottom: float


class PdfPageOut(BaseModel):
    page: int
    width: float
    height: float


class DocumentLayoutOut(BaseModel):
    pages: list[PdfPageOut]
    finding_rects: dict[str, list[PdfRectOut]]
    annotation_rects: dict[str, list[PdfRectOut]]
    clause_rects: dict[str, list[PdfRectOut]]


# ---------------------------------------------------------------------------
# PDF text anchoring helpers
# ---------------------------------------------------------------------------


_TOKEN_RE = re.compile(r"[A-Za-z0-9]+(?:['-][A-Za-z0-9]+)?")
# Matches ASCII quotes plus Unicode curly/smart quotes (U+2018-U+201D)
_QUOTE_RE = re.compile(
    "[''""\x27\x22]"
    "([^''""\x27\x22]{4,100})"
    "[''""\x27\x22]"
)


def _tokens(text: str) -> list[str]:
    return [m.group(0).lower() for m in _TOKEN_RE.finditer(text or "")]


def _escape_pdf_text(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _words_to_line_rects(page_num: int, words: list[dict]) -> list[PdfRectOut]:
    if not words:
        return []
    lines: list[list[dict]] = []
    for word in words:
        placed = False
        mid = (float(word["top"]) + float(word["bottom"])) / 2
        for line in lines:
            line_mid = (float(line[0]["top"]) + float(line[0]["bottom"])) / 2
            if abs(mid - line_mid) <= 3:
                line.append(word)
                placed = True
                break
        if not placed:
            lines.append([word])
    rects: list[PdfRectOut] = []
    for line in lines:
        rects.append(
            PdfRectOut(
                page=page_num,
                x0=min(float(w["x0"]) for w in line),
                top=min(float(w["top"]) for w in line),
                x1=max(float(w["x1"]) for w in line),
                bottom=max(float(w["bottom"]) for w in line),
            )
        )
    return rects


def _find_text_rects(page_num: int, words: list[dict], needle: str, *, max_tokens: int = 80) -> list[PdfRectOut]:
    needle_tokens = _tokens(needle)
    if not needle_tokens or len(needle_tokens) > max_tokens:
        return []
    word_tokens = [_tokens(str(w.get("text", "")))[0] if _tokens(str(w.get("text", ""))) else "" for w in words]
    if len(needle_tokens) > len(word_tokens):
        return []
    for idx in range(0, len(word_tokens) - len(needle_tokens) + 1):
        if word_tokens[idx:idx + len(needle_tokens)] == needle_tokens:
            return _words_to_line_rects(page_num, words[idx:idx + len(needle_tokens)])
    return []


def _candidate_needles_from_finding(finding) -> list[str]:
    candidates: list[str] = []
    source_text = " ".join(
        [
            getattr(finding, "description", "") or "",
            getattr(finding, "recommendation_detail", "") or "",
            getattr(finding, "recommended_change", "") or "",
        ]
    )
    # Use global _QUOTE_RE that handles ASCII and Unicode smart quotes
    for quoted in _QUOTE_RE.findall(source_text):
        stripped = quoted.strip()
        if stripped:
            candidates.append(stripped)

    def _add_evidence_needles(text: str) -> None:
        toks = _tokens(text)
        if not toks:
            return
        # Prefer a short precise phrase; take first 10 tokens for highlight precision
        needle = " ".join(text.split()[:min(10, len(toks))])
        candidates.append(needle)
        # Also try the full text if it's short enough for an exact match
        if len(toks) <= 60 and len(toks) > 10:
            candidates.append(text)

    for ev in getattr(finding, "contract_evidence", []) or []:
        text = (getattr(ev, "text", "") or "").strip()
        if text:
            _add_evidence_needles(text)
    for ev in getattr(finding, "evidence", []) or []:
        text = (getattr(ev, "normalized_text", "") or "").strip()
        if text:
            _add_evidence_needles(text)

    seen: set[str] = set()
    unique: list[str] = []
    for candidate in candidates:
        key = " ".join(_tokens(candidate))
        if key and key not in seen:
            seen.add(key)
            unique.append(candidate)
    return unique


def _rects_to_pdf_content(rects: list[PdfRectOut], page_height: float, rgb: tuple[float, float, float], alpha_hint: float = 0.22) -> str:
    # pypdf cannot add transparency without a heavier graphics-state setup. Use pale fills.
    lines = []
    r, g, b = rgb
    for rect in rects:
        x = rect.x0
        y = page_height - rect.bottom
        w = max(0.0, rect.x1 - rect.x0)
        h = max(0.0, rect.bottom - rect.top)
        if w <= 0 or h <= 0:
            continue
        lines.append(f"q {r:.3f} {g:.3f} {b:.3f} rg {x:.2f} {y:.2f} {w:.2f} {h:.2f} re f Q")
    return "\n".join(lines)


def _plain_text_from_edit(edit: dict | str) -> str:
    if isinstance(edit, str):
        return edit
    if not isinstance(edit, dict):
        return ""
    value = edit.get("plain_text") or edit.get("text") or ""
    return re.sub(r"<[^>]+>", "", str(value)).replace("&nbsp;", " ").strip()


def _safe_rich_text_to_plain(text: str) -> str:
    text = html.unescape(text or "")
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.I)
    text = re.sub(r"</(p|div)>", "\n", text, flags=re.I)
    text = re.sub(r"<[^>]+>", "", text)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def _build_document_blocks(
    clauses: list,
    edits: dict,
    original: bool = False,
) -> list[tuple[str, str]]:
    """Return (style, text) pairs for clean final contract export with no redlines."""
    blocks: list[tuple[str, str]] = []
    for clause in clauses:
        raw_edit = None if original else edits.get(clause.clause_uid)
        if raw_edit and isinstance(raw_edit, dict):
            full_text = raw_edit.get("plain_text") or _safe_rich_text_to_plain(raw_edit.get("text", ""))
        else:
            full_text = clause.normalized_text
        full_text = re.sub(r"\*{1,2}([^*]+)\*{1,2}", r"\1", full_text).strip()
        if not full_text:
            continue

        if len(clause.normalized_text) > 300:
            sub_blocks = _split_contract_blob(clause.normalized_text, clause.clause_uid, clause.page_number)
            parent_edit_applied = False
            for sub_uid, sub_style, sub_text in sub_blocks:
                if original:
                    display_text = sub_text
                else:
                    is_sub_uid = sub_uid != clause.clause_uid
                    sub_edit = edits.get(sub_uid) if is_sub_uid else None
                    if sub_edit is None and raw_edit and isinstance(raw_edit, dict) and not parent_edit_applied and sub_style in ("body", "list_item"):
                        sub_edit = raw_edit
                        parent_edit_applied = True
                    if sub_edit and isinstance(sub_edit, dict):
                        display_text = sub_edit.get("plain_text") or _safe_rich_text_to_plain(sub_edit.get("text", "")) or sub_text
                    else:
                        display_text = sub_text
                display_text = re.sub(r"\*{1,2}([^*]+)\*{1,2}", r"\1", display_text).strip()
                if display_text:
                    blocks.append((sub_style, display_text))
        else:
            style_key = _infer_block_style(clause.normalized_text)
            blocks.append((style_key, full_text))
    return blocks


def _build_clean_pdf(blocks: list[tuple[str, str]]) -> io.BytesIO:
    """Generate a clean contract PDF from (style, text) blocks using reportlab."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    except ImportError:
        raise HTTPException(status_code=501, detail="reportlab not installed; run: pip install reportlab")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        leftMargin=1.25 * inch,
        rightMargin=1.25 * inch,
        topMargin=inch,
        bottomMargin=inch,
    )
    _h1 = ParagraphStyle("h1", fontName="Times-Bold", fontSize=16, spaceAfter=6, alignment=TA_CENTER, leading=20)
    _h2 = ParagraphStyle("h2", fontName="Times-Bold", fontSize=12, spaceBefore=12, spaceAfter=4, alignment=TA_LEFT, leading=16)
    _h3 = ParagraphStyle("h3", fontName="Times-Bold", fontSize=11, spaceBefore=8, spaceAfter=2, alignment=TA_LEFT, leading=15)
    _sub = ParagraphStyle("sub", fontName="Times-Roman", fontSize=12, spaceAfter=12, alignment=TA_CENTER, leading=16)
    _body = ParagraphStyle("body", fontName="Times-Roman", fontSize=11, spaceAfter=8, alignment=TA_JUSTIFY, leading=14)
    style_map = {"heading1": _h1, "heading2": _h2, "heading3": _h3, "subtitle": _sub, "body": _body, "list_item": _body}

    story = []
    for style_key, text in blocks:
        if not text.strip():
            continue
        safe = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        story.append(Paragraph(safe, style_map.get(style_key, _body)))

    if not story:
        story.append(Paragraph("(empty contract)", _body))
    doc.build(story)
    buf.seek(0)
    return buf


def _find_pending_sub_uid(sub_blocks: list[tuple[str, str, str]], finding) -> str | None:
    """Return sub-block UID that best matches the finding's evidence text; falls back to first body block."""
    evidence_text = ""
    for ev in (getattr(finding, "contract_evidence", None) or []):
        t = (getattr(ev, "text", "") or "").strip()
        if t:
            evidence_text = t
            break
    if evidence_text:
        needle = evidence_text[:60].lower()
        for sub_uid, sub_style, sub_text in sub_blocks:
            if sub_style in ("body", "list_item") and needle in sub_text.lower():
                return sub_uid
    for sub_uid, sub_style, sub_text in sub_blocks:
        if sub_style in ("body", "list_item"):
            return sub_uid
    return None


# ---------------------------------------------------------------------------
# pdf2docx-based high-fidelity export helpers
# ---------------------------------------------------------------------------


def _norm_text(text: str) -> str:
    return " ".join(text.lower().split())


def _extract_docx_blocks(docx_path: str) -> list[dict]:
    """Parse a python-docx file and return one entry per non-empty paragraph.

    Each entry::

        {
            "text":  str,                           # plain paragraph text
            "style": "heading1"|...|"body",
            "marks": [{"type": "bold"|"italic", "from_pos": int, "to_pos": int}],
        }
    """
    try:
        from docx import Document as _DocxDoc
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return []

    doc = _DocxDoc(docx_path)
    result: list[dict] = []

    for para in doc.paragraphs:
        text = para.text
        if not text.strip():
            continue

        # ---- collect marks and detect dominant bold/size ----
        marks: list[dict] = []
        offset = 0
        any_bold = False
        max_pt: float = 0.0
        # Use paragraph-level style name as fallback hint
        para_style_name = (para.style.name or "").lower() if para.style else ""

        for run in para.runs:
            run_len = len(run.text)
            if run_len == 0:
                continue
            is_bold = bool(run.bold)
            is_italic = bool(run.italic)
            if is_bold:
                any_bold = True
                marks.append({"type": "bold", "from_pos": offset, "to_pos": offset + run_len})
            if is_italic:
                marks.append({"type": "italic", "from_pos": offset, "to_pos": offset + run_len})
            # Track largest font size across runs
            if run.font.size:
                pt = run.font.size.pt
                if pt and pt > max_pt:
                    max_pt = pt
            offset += run_len

        # ---- infer style ----
        is_centered = False
        try:
            is_centered = para.alignment in (WD_ALIGN_PARAGRAPH.CENTER,)
        except Exception:
            pass

        if "heading 1" in para_style_name:
            style = "heading1"
        elif "heading 2" in para_style_name:
            style = "heading2"
        elif "heading 3" in para_style_name:
            style = "heading3"
        elif any_bold and max_pt >= 14:
            style = "heading1"
        elif any_bold and max_pt >= 12:
            style = "heading2"
        elif any_bold:
            style = "heading3"
        elif is_centered and max_pt >= 12:
            style = "subtitle"
        else:
            style = "body"

        result.append({"text": text, "style": style, "marks": marks})

    return result


def _replace_para_text(para, new_text: str) -> None:
    """Replace all runs in a paragraph with one clean run, keeping font name/size."""
    from docx.oxml.ns import qn
    p = para._p
    font_name = None
    font_size = None
    for run in para.runs:
        if not font_name and run.font.name:
            font_name = run.font.name
        if not font_size and run.font.size:
            font_size = run.font.size
        if font_name and font_size:
            break
    for r in list(p.findall(qn("w:r"))):
        p.remove(r)
    for hyp in list(p.findall(qn("w:hyperlink"))):
        p.remove(hyp)
    new_run = para.add_run(new_text)
    if font_name:
        new_run.font.name = font_name
    if font_size:
        new_run.font.size = font_size


def _build_docx_edit_map(clauses: list, edits: dict) -> list[tuple[str, str, str]]:
    """Return [(norm_prefix_30, full_norm_original, replacement_text)] for in-place DOCX paragraph matching."""
    clause_by_uid = {c.clause_uid: c for c in clauses}
    result: list[tuple[str, str, str]] = []
    for clause_uid, raw_edit in edits.items():
        replacement = _plain_text_from_edit(raw_edit)
        replacement = re.sub(r"\*{1,2}([^*]+)\*{1,2}", r"\1", replacement).strip()
        if not replacement:
            continue
        parent_uid = clause_uid.rsplit("__p", 1)[0] if "__p" in clause_uid else clause_uid
        clause = clause_by_uid.get(parent_uid)
        if not clause:
            continue
        if "__p" in clause_uid:
            sub_blocks = _split_contract_blob(clause.normalized_text, parent_uid, clause.page_number)
            for sub_uid, _, sub_text in sub_blocks:
                if sub_uid == clause_uid:
                    full_norm = _norm_text(sub_text)
                    result.append((full_norm[:30], full_norm, replacement))
                    break
        else:
            full_norm = _norm_text(clause.normalized_text)
            result.append((full_norm[:30], full_norm, replacement))
    return result


def _apply_edits_to_docx(docx_path: str, clauses: list, edits: dict):
    """Load a pdf2docx DOCX, apply accepted edits in-place, return modified Document."""
    from docx import Document as _DocxDoc
    doc = _DocxDoc(docx_path)
    if not edits:
        return doc
    edit_map = _build_docx_edit_map(clauses, edits)
    if not edit_map:
        return doc
    applied: set[int] = set()
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        para_norm = _norm_text(text)
        for idx, (orig_prefix, full_norm_orig, replacement) in enumerate(edit_map):
            if idx in applied:
                continue
            check = min(len(orig_prefix), 25)
            matched = False
            # Strategy 1: exact-prefix match
            if para_norm[:check] == orig_prefix[:check]:
                matched = True
            # Strategy 2: orig_prefix appears anywhere in paragraph
            elif orig_prefix and orig_prefix in para_norm:
                matched = True
            # Strategy 3: paragraph text appears as substring of original (for split paragraphs)
            elif len(para_norm) > 25 and para_norm in full_norm_orig:
                matched = True
            if matched:
                _replace_para_text(para, replacement)
                applied.add(idx)
                break
    logger.info("_apply_edits_to_docx: edits=%d paragraphs=%d applied=%d", len(edit_map), len(doc.paragraphs), len(applied))
    return doc


def _docx_to_pdf_buf(doc) -> "io.BytesIO | None":
    """Convert python-docx Document → PDF.

    Tier 1 (primary): LibreOffice headless — most reliable, deterministic, cross-platform.
    Tier 2 (fallback): docx2pdf (Word COM on Windows) — used only when LibreOffice is missing.
    Returns None if both tiers fail.
    """
    import tempfile
    import shutil
    import subprocess
    from pathlib import Path as _Path

    # Tier 1: LibreOffice headless
    soffice_bin = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice_bin:
        # Common Windows install locations not on PATH
        for candidate in (
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ):
            if _Path(candidate).exists():
                soffice_bin = candidate
                break

    if soffice_bin:
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                docx_path = _Path(tmpdir) / "contract.docx"
                doc.save(str(docx_path))
                result = subprocess.run(
                    [soffice_bin, "--headless", "--convert-to", "pdf", "--outdir", tmpdir, str(docx_path)],
                    capture_output=True,
                    timeout=120,
                )
                pdf_path = _Path(tmpdir) / "contract.pdf"
                if result.returncode == 0 and pdf_path.exists() and pdf_path.stat().st_size > 0:
                    logger.info("docx→pdf via libreoffice OK (%s)", soffice_bin)
                    return io.BytesIO(pdf_path.read_bytes())
                logger.warning(
                    "libreoffice convert failed: rc=%s stdout=%r stderr=%r",
                    result.returncode, result.stdout[-200:], result.stderr[-200:],
                )
        except Exception as e:
            logger.warning("libreoffice subprocess raised: %s: %s", type(e).__name__, e)
    else:
        logger.info("libreoffice not found on PATH or standard install dirs; trying docx2pdf fallback")

    # Tier 2: docx2pdf (Word COM)
    try:
        import docx2pdf
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = _Path(tmpdir) / "contract.docx"
            pdf_path = _Path(tmpdir) / "contract.pdf"
            doc.save(str(docx_path))
            docx2pdf.convert(str(docx_path), str(pdf_path))
            if pdf_path.exists() and pdf_path.stat().st_size > 0:
                logger.info("docx→pdf via docx2pdf OK")
                return io.BytesIO(pdf_path.read_bytes())
    except Exception as e:
        logger.warning("docx2pdf failed: %s: %s", type(e).__name__, e)
    return None


def _build_modified_docx(run, clauses: list, edits: dict):
    """Return a python-docx Document for the modified contract (single source of truth for DOCX + PDF export).

    Strategy:
      1. If pdf2docx cache exists → load it and apply edits in-place (preserves original PDF formatting).
      2. Otherwise → build from scratch using parsed clauses + style heuristics.
    """
    if run.storage_path:
        cache = Path(run.storage_path).with_suffix(".original.docx")
        if cache.exists():
            try:
                return _apply_edits_to_docx(str(cache), list(clauses), edits)
            except Exception:
                pass

    # Fallback: build from parsed clauses
    from docx import Document as _DocxDoc
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn as _qn

    doc = _DocxDoc()
    try:
        settings_el = doc.settings.element
        for tag in ("w:trackChanges", "w:revisionView"):
            for el in settings_el.findall(_qn(tag)):
                settings_el.remove(el)
    except Exception:
        pass
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    def _apply_style(p, style_key: str, text: str) -> None:
        pf = p.paragraph_format
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        if style_key == "heading1":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r.font.bold = True; r.font.size = Pt(16); pf.space_after = Pt(6)
        elif style_key == "subtitle":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r.font.size = Pt(12); pf.space_after = Pt(12)
        elif style_key == "heading2":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r.font.bold = True; r.font.size = Pt(12); pf.space_before = Pt(12); pf.space_after = Pt(4)
        elif style_key == "heading3":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r.font.bold = True; r.font.size = Pt(11); pf.space_before = Pt(8); pf.space_after = Pt(2)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r.font.size = Pt(11); pf.space_after = Pt(8)

    blocks = _build_document_blocks(list(clauses), edits, original=False)
    for style_key, text in blocks:
        if not text.strip():
            continue
        _apply_style(doc.add_paragraph(), style_key, text)
    return doc


def _pdf_from_docx_doc(doc) -> io.BytesIO:
    """Build a reportlab PDF from a python-docx Document, detecting bold/size for headings."""
    blocks: list[tuple[str, str]] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        is_bold = any(r.bold for r in para.runs if r.text.strip())
        font_pt: float | None = None
        for r in para.runs:
            if r.font.size:
                font_pt = r.font.size.pt
                break
        if is_bold and font_pt and font_pt >= 14:
            style_key = "heading1"
        elif is_bold:
            style_key = "heading2"
        else:
            style_key = "body"
        blocks.append((style_key, text))
    return _build_clean_pdf(blocks)


# ---------------------------------------------------------------------------
# Clause listing
# ---------------------------------------------------------------------------


@router.get("/runs/{run_id}/clauses", response_model=list[ClauseOut])
async def list_clauses(run_id: str, db: DbSession) -> list[ClauseOut]:
    result = await db.execute(
        select(ParsedClauseRecord)
        .where(ParsedClauseRecord.run_id == run_id)
        .order_by(ParsedClauseRecord.order_index.asc().nullslast(), ParsedClauseRecord.page_number.asc())
    )
    clauses = result.scalars().all()
    return [
        ClauseOut(
            clause_uid=c.clause_uid,
            page_number=c.page_number,
            normalized_text=c.normalized_text,
            bbox=[c.bbox_x0 or 0.0, c.bbox_y0 or 0.0, c.bbox_x1 or 0.0, c.bbox_y1 or 0.0],
        )
        for c in clauses
    ]


# ---------------------------------------------------------------------------
# Contract-level clause edits (stored on the run)
# ---------------------------------------------------------------------------


@router.get("/runs/{run_id}/contract-edits")
async def get_contract_edits(run_id: str, db: DbSession) -> dict:
    run = (await db.execute(select(RunRecord).where(RunRecord.id == run_id))).scalar_one_or_none()
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")
    return run.contract_edits or {}


@router.get("/runs/{run_id}/document-layout", response_model=DocumentLayoutOut)
async def get_document_layout(run_id: str, db: DbSession) -> DocumentLayoutOut:
    """Resolve pages, clause boxes, and exact finding/comment highlight rectangles."""
    run = (await db.execute(select(RunRecord).where(RunRecord.id == run_id))).scalar_one_or_none()
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")

    path = Path(run.storage_path)
    if not path.exists():
        raise HTTPException(status_code=404, detail="File not found on disk")

    clauses = (await db.execute(
        select(ParsedClauseRecord)
        .where(ParsedClauseRecord.run_id == run_id)
        .order_by(ParsedClauseRecord.order_index.asc().nullslast(), ParsedClauseRecord.page_number.asc(), ParsedClauseRecord.bbox_y0.asc())
    )).scalars().all()

    annotations = (await db.execute(
        select(DocumentAnnotationRecord).where(
            DocumentAnnotationRecord.run_id == run_id,
            DocumentAnnotationRecord.status.not_in(["accepted", "dismissed", "deleted"]),
        )
    )).scalars().all()

    from app.backend.services.run_service import _load_full_findings
    import pdfplumber

    pages: list[PdfPageOut] = []
    words_by_page: dict[int, list[dict]] = {}
    with pdfplumber.open(str(path)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            pages.append(PdfPageOut(page=page_num, width=float(page.width), height=float(page.height)))
            words_by_page[page_num] = page.extract_words() or []

    clause_rects: dict[str, list[PdfRectOut]] = {}
    for clause in clauses:
        page_words = words_by_page.get(clause.page_number, [])
        rects = _find_text_rects(clause.page_number, page_words, clause.normalized_text, max_tokens=180)
        if rects:
            clause_rects[clause.clause_uid] = rects

    finding_rects: dict[str, list[PdfRectOut]] = {}
    for finding in await _load_full_findings(db, run_id):
        page_num = 1
        if getattr(finding, "contract_evidence", None):
            page_num = finding.contract_evidence[0].page
        elif getattr(finding, "evidence", None):
            page_num = finding.evidence[0].page
        page_words = words_by_page.get(page_num, [])
        for needle in _candidate_needles_from_finding(finding):
            rects = _find_text_rects(page_num, page_words, needle, max_tokens=80)
            if rects:
                finding_rects[finding.finding_id] = rects
                break

    annotation_rects: dict[str, list[PdfRectOut]] = {}
    for ann in annotations:
        if not ann.selected_text:
            continue
        page_num = ann.page_number or 1
        rects = _find_text_rects(page_num, words_by_page.get(page_num, []), ann.selected_text, max_tokens=80)
        if not rects:
            for pn, words in words_by_page.items():
                if pn == page_num:
                    continue
                rects = _find_text_rects(pn, words, ann.selected_text, max_tokens=80)
                if rects:
                    break
        if rects:
            annotation_rects[ann.id] = rects

    return DocumentLayoutOut(
        pages=pages,
        finding_rects=finding_rects,
        annotation_rects=annotation_rects,
        clause_rects=clause_rects,
    )


@router.get("/runs/{run_id}/export-edited")
async def export_edited_contract(run_id: str, db: DbSession) -> StreamingResponse:
    """Export the modified contract as PDF.

    Always goes DOCX → PDF: builds the same modified DOCX that /export-docx returns,
    then converts it to PDF via docx2pdf (Word/LibreOffice). This guarantees the
    downloaded PDF matches the downloaded DOCX byte-for-byte in layout.
    """
    run = (await db.execute(select(RunRecord).where(RunRecord.id == run_id))).scalar_one_or_none()
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")

    clauses = (await db.execute(
        select(ParsedClauseRecord)
        .where(ParsedClauseRecord.run_id == run_id)
        .order_by(ParsedClauseRecord.order_index.asc().nullslast(), ParsedClauseRecord.page_number.asc(), ParsedClauseRecord.bbox_y0.asc())
    )).scalars().all()

    edits: dict = run.contract_edits or {}
    base = (run.original_filename or "contract").rsplit(".", 1)[0]

    # Prime the pdf2docx cache so the modified DOCX has full original formatting
    cache_path = None
    if run.storage_path:
        pdf_path = Path(run.storage_path)
        if pdf_path.exists():
            cache_path = pdf_path.with_suffix(".original.docx")
            if not cache_path.exists():
                try:
                    from pdf2docx import Converter as _Cv
                    cv = _Cv(str(pdf_path))
                    cv.convert(str(cache_path), start=0, end=None)
                    cv.close()
                except Exception:
                    pass

    logger.info("export-edited: run=%s cache=%s edits=%d", run_id, cache_path.exists() if cache_path else "no-storage", len(edits))

    # Step 1: build the same modified DOCX that /export-docx returns
    modified_doc = _build_modified_docx(run, list(clauses), edits)

    # Step 2: convert that DOCX → PDF using docx2pdf
    pdf_buf = _docx_to_pdf_buf(modified_doc)
    logger.info("export-edited: docx2pdf result=%s", "ok" if pdf_buf else "failed")
    if pdf_buf is None:
        raise HTTPException(
            status_code=503,
            detail=(
                "PDF export requires LibreOffice (recommended, free) or Microsoft Word installed for DOCX→PDF conversion. "
                "Install LibreOffice from https://www.libreoffice.org/download/ , or download the DOCX export instead."
            ),
        )
    return StreamingResponse(
        pdf_buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{base}-edited.pdf"'},
    )


@router.get("/runs/{run_id}/export-edited-legacy")
async def export_edited_contract_legacy(run_id: str, db: DbSession) -> StreamingResponse:
    """Legacy: export original PDF with redline overlays. Kept for debugging only."""
    run = (await db.execute(select(RunRecord).where(RunRecord.id == run_id))).scalar_one_or_none()
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")

    path = Path(run.storage_path)
    if not path.exists():
        raise HTTPException(status_code=404, detail="File not found on disk")

    clauses = (await db.execute(
        select(ParsedClauseRecord)
        .where(ParsedClauseRecord.run_id == run_id)
        .order_by(ParsedClauseRecord.order_index.asc().nullslast(), ParsedClauseRecord.page_number.asc(), ParsedClauseRecord.bbox_y0.asc())
    )).scalars().all()

    annotations_rows = (await db.execute(
        select(DocumentAnnotationRecord)
        .where(
            DocumentAnnotationRecord.run_id == run_id,
            DocumentAnnotationRecord.status.not_in(["dismissed", "deleted"]),
        )
        .order_by(DocumentAnnotationRecord.created_at.asc())
    )).scalars().all()

    edits: dict = run.contract_edits or {}
    from app.backend.services.run_service import _load_full_findings
    from pypdf import PdfReader, PdfWriter
    from pypdf.generic import ArrayObject, DecodedStreamObject, DictionaryObject, NameObject
    import pdfplumber

    pages: dict[int, dict] = {}
    with pdfplumber.open(str(path)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            pages[page_num] = {
                "width": float(page.width),
                "height": float(page.height),
                "words": page.extract_words() or [],
            }

    clause_by_uid = {c.clause_uid: c for c in clauses}
    content_by_page: dict[int, list[str]] = {}

    # Annotation highlights go BEFORE the original content stream (prepend) so
    # the original text renders on top and remains readable.
    highlights_by_page: dict[int, list[str]] = {}
    # Edit overlays go AFTER the original content stream (append) so they
    # white-out original text and show the replacement on top.
    edits_by_page: dict[int, list[str]] = {}

    for ann in annotations_rows:
        if not ann.selected_text:
            continue
        page_num = ann.page_number or 1
        page_info = pages.get(page_num)
        rects = _find_text_rects(page_num, (page_info or {}).get("words", []), ann.selected_text, max_tokens=80) if page_info else []
        actual_page = page_num
        if not rects:
            for pn, pi in pages.items():
                if pn == page_num:
                    continue
                rects = _find_text_rects(pn, pi["words"], ann.selected_text, max_tokens=80)
                if rects:
                    actual_page = pn
                    page_info = pi
                    break
        if rects and page_info:
            highlights_by_page.setdefault(actual_page, []).append(
                _rects_to_pdf_content(rects, page_info["height"], (0.82, 0.89, 1.0))
            )

    for clause_uid, raw_edit in edits.items():
        clause = clause_by_uid.get(clause_uid)
        # Sub-block UID (e.g. parent__p2): find parent clause then reconstruct sub-block text
        original_needle_text: str | None = None
        if not clause and "__p" in clause_uid:
            parent_uid = clause_uid.rsplit("__p", 1)[0]
            clause = clause_by_uid.get(parent_uid)
            if clause:
                sub_blocks = _split_contract_blob(clause.normalized_text, parent_uid, clause.page_number)
                for sub_uid, _sub_style, sub_text in sub_blocks:
                    if sub_uid == clause_uid:
                        original_needle_text = sub_text
                        break

        if not clause:
            continue

        if original_needle_text is None:
            original_needle_text = clause.normalized_text

        page_num = int((raw_edit or {}).get("page") or clause.page_number) if isinstance(raw_edit, dict) else clause.page_number
        page_info = pages.get(page_num)
        if not page_info:
            continue
        edit_text = _plain_text_from_edit(raw_edit)
        if not edit_text:
            continue
        anchor_text = (raw_edit or {}).get("anchor_text", "") if isinstance(raw_edit, dict) else ""
        rects: list[PdfRectOut] = []
        if isinstance(raw_edit, dict) and raw_edit.get("rects"):
            rects = [PdfRectOut(**r) for r in raw_edit["rects"]]
        if not rects:
            rects = _find_text_rects(page_num, page_info["words"], original_needle_text, max_tokens=180)
        if not rects and anchor_text:
            rects = _find_text_rects(page_num, page_info["words"], anchor_text, max_tokens=180)
        # For long text the full needle exceeds max_tokens — locate via progressively shorter prefix
        if not rects:
            needle_words = original_needle_text.split()
            for prefix_len in (40, 20, 10):
                if len(needle_words) >= prefix_len:
                    prefix = " ".join(needle_words[:prefix_len])
                    rects = _find_text_rects(page_num, page_info["words"], prefix, max_tokens=prefix_len + 5)
                    if rects:
                        break
        # Try across all pages when the stored page_num is wrong
        if not rects:
            needle_words = original_needle_text.split()
            for pn, pi in pages.items():
                if pn == page_num:
                    continue
                for prefix_len in (40, 20):
                    if len(needle_words) >= prefix_len:
                        prefix = " ".join(needle_words[:prefix_len])
                        rects = _find_text_rects(pn, pi["words"], prefix, max_tokens=prefix_len + 5)
                        if rects:
                            page_num = pn
                            page_info = pi
                            break
                if rects:
                    break

        # Skip edits we cannot locate in the PDF — avoid random placement
        if not rects:
            continue

        x0 = min(r.x0 for r in rects)
        top = min(r.top for r in rects)
        x1 = max(r.x1 for r in rects)
        bottom = max(r.bottom for r in rects)
        orig_h = max(14.0, bottom - top)
        orig_w = max(120.0, x1 - x0)

        # --- Redline: strikethrough over original text area ---
        # Draw horizontal lines at roughly 40% height of each "line" in the bounding box
        line_height_est = 12.0
        num_orig_lines = max(1, round(orig_h / line_height_est))
        overlay: list[str] = []
        for li in range(num_orig_lines):
            line_mid_top = top + (li + 0.55) * (orig_h / num_orig_lines)
            strike_y = page_info["height"] - line_mid_top
            overlay.append(
                f"q 0.8 0.1 0.1 RG 0.8 w {x0:.2f} {strike_y:.2f} m {x1:.2f} {strike_y:.2f} l S Q"
            )

        # --- Replacement box immediately below original ---
        gap = 4.0
        font_size = 9.5
        repl_line_h = 12.0
        max_chars = max(30, int(orig_w / 5.5))
        words_list = edit_text.split()
        lines: list[str] = []
        current = ""
        for word in words_list:
            candidate = f"{current} {word}".strip()
            if len(candidate) > max_chars:
                if current:
                    lines.append(current)
                current = word
            else:
                current = candidate
        if current:
            lines.append(current)

        repl_h = len(lines) * repl_line_h + 14.0
        repl_top = bottom + gap
        repl_y_pdf = page_info["height"] - repl_top - repl_h

        # Light green-tinted background for the replacement box
        overlay.append(f"q 0.94 0.99 0.95 rg {x0:.2f} {repl_y_pdf:.2f} {orig_w:.2f} {repl_h:.2f} re f Q")
        # Blue left-edge bar
        overlay.append(f"q 0.149 0.392 0.863 rg {x0:.2f} {repl_y_pdf:.2f} 2.5 {repl_h:.2f} re f Q")
        # "Replacement" label
        label_y = page_info["height"] - repl_top - 8.0
        overlay.append(
            f"BT /FvHelvetica 7 Tf 0.149 0.392 0.863 rg "
            f"1 0 0 1 {x0 + 6:.2f} {label_y:.2f} Tm (REPLACEMENT) Tj ET"
        )
        # Replacement text lines
        for idx, line in enumerate(lines[:25]):
            line_y = page_info["height"] - repl_top - 14.0 - (idx * repl_line_h)
            if line_y < 10:
                break
            overlay.append(
                f"BT /FvHelvetica {font_size} Tf 0.067 0.247 0.067 rg "
                f"1 0 0 1 {x0 + 6:.2f} {line_y:.2f} Tm ({_escape_pdf_text(line)}) Tj ET"
            )
        edits_by_page.setdefault(page_num, []).append("\n".join(overlay))

    reader = PdfReader(str(path))
    writer = PdfWriter()

    def _setup_font(out_page):
        resources_obj = out_page.get("/Resources")
        resources = resources_obj.get_object() if hasattr(resources_obj, "get_object") else (resources_obj or DictionaryObject())
        fonts_obj = resources.get("/Font") if resources else None
        fonts = fonts_obj.get_object() if hasattr(fonts_obj, "get_object") else (fonts_obj or DictionaryObject())
        fonts[NameObject("/FvHelvetica")] = DictionaryObject({
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        })
        resources[NameObject("/Font")] = fonts
        out_page[NameObject("/Resources")] = resources

    def _prepend_stream(out_page, data: str):
        """Insert stream BEFORE original content so original text renders on top."""
        obj = DecodedStreamObject()
        obj.set_data(data.encode("utf-8"))
        ref = writer._add_object(obj)
        contents = out_page.get("/Contents")
        if contents is None:
            out_page[NameObject("/Contents")] = ref
        elif isinstance(contents, ArrayObject):
            contents.insert(0, ref)
        else:
            out_page[NameObject("/Contents")] = ArrayObject([ref, contents])

    def _append_stream(out_page, data: str):
        """Append stream AFTER original content so it paints over the original."""
        obj = DecodedStreamObject()
        obj.set_data(data.encode("utf-8"))
        ref = writer._add_object(obj)
        contents = out_page.get("/Contents")
        if contents is None:
            out_page[NameObject("/Contents")] = ref
        elif isinstance(contents, ArrayObject):
            contents.append(ref)
        else:
            out_page[NameObject("/Contents")] = ArrayObject([contents, ref])

    for index, page in enumerate(reader.pages, start=1):
        writer.add_page(page)
        hl = highlights_by_page.get(index, [])
        ed = edits_by_page.get(index, [])
        if not hl and not ed:
            continue
        out_page = writer.pages[index - 1]
        _setup_font(out_page)
        if hl:
            _prepend_stream(out_page, "\n".join(hl))
        if ed:
            _append_stream(out_page, "\n".join(ed))

    buf = io.BytesIO()
    writer.write(buf)
    buf.seek(0)
    base = (run.original_filename or "contract").rsplit(".", 1)[0]
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{base}-edited.pdf"'},
    )


@router.put("/runs/{run_id}/contract-edits/{clause_uid}", status_code=status.HTTP_200_OK)
async def save_clause_edit(
    run_id: str,
    clause_uid: str,
    body: ClauseEditBody,
    db: DbSession,
    token: dict = Depends(require_auth),
) -> dict:
    run = (await db.execute(select(RunRecord).where(RunRecord.id == run_id))).scalar_one_or_none()
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")
    edits: dict = dict(run.contract_edits or {})
    plain_text = body.plain_text or _safe_rich_text_to_plain(body.text)
    edit_payload = {
        "schema_version": 2,
        "text": body.text,
        "plain_text": plain_text,
        "rich_text": body.rich_text or [],
        "page": body.page,
        "rects": body.rects or [],
        "anchor_text": body.anchor_text,
        "edited_at": datetime.utcnow().isoformat(),
    }
    edits[clause_uid] = edit_payload
    run.contract_edits = edits
    flag_modified(run, "contract_edits")
    await db.commit()
    return {"clause_uid": clause_uid, **edit_payload}


# ---------------------------------------------------------------------------
# Finding accept / dismiss
# ---------------------------------------------------------------------------


@router.post("/runs/{run_id}/findings/{finding_id}/accept", status_code=status.HTTP_200_OK)
async def accept_finding(
    run_id: str,
    finding_id: str,
    db: DbSession,
    body: AcceptFindingBody | None = Body(default=None),
    token: dict = Depends(require_auth),
) -> dict:
    finding = (await db.execute(
        select(FindingRecord).where(FindingRecord.id == finding_id, FindingRecord.run_id == run_id)
    )).scalar_one_or_none()
    full_finding = None
    if not finding:
        from app.backend.services.run_service import _load_full_findings
        full_finding = next((f for f in await _load_full_findings(db, run_id) if f.finding_id == finding_id), None)
        if full_finding is None:
            raise HTTPException(status_code=404, detail="Finding not found")

    # Determine replacement text: explicit override -> finding.recommended_change -> None
    body = body or AcceptFindingBody()
    replacement_text = body.custom_text or (
        finding.recommended_change if finding else getattr(full_finding, "recommended_change", None)
    )
    clause_uid = finding.clause_uid if finding else getattr(full_finding, "clause_uid", None)

    applied_text: str | None = None
    if replacement_text and clause_uid:
        run = (await db.execute(select(RunRecord).where(RunRecord.id == run_id))).scalar_one_or_none()
        if run:
            # Determine the correct UID to store the edit under.
            # For large-blob clauses the viewer splits into sub-blocks (synthetic __pN UIDs).
            # We match the finding's evidence text against those sub-blocks so the edit is
            # stored under the sub-block UID, keeping it consistent with manual editor edits.
            target_uid = clause_uid
            clause_rec = (await db.execute(
                select(ParsedClauseRecord).where(
                    ParsedClauseRecord.run_id == run_id,
                    ParsedClauseRecord.clause_uid == clause_uid,
                )
            )).scalar_one_or_none()
            if clause_rec and len(clause_rec.normalized_text or "") > 300:
                evidence_text = ""
                ev_list = (getattr(finding, "contract_evidence", None) or
                           getattr(full_finding, "contract_evidence", None) or [])
                for ev in ev_list:
                    excerpt = (ev or {}).get("excerpt") or (ev or {}).get("text") or ""
                    if excerpt:
                        evidence_text = excerpt
                        break
                if not evidence_text:
                    ev_list2 = (getattr(finding, "evidence", None) or
                                getattr(full_finding, "evidence", None) or [])
                    for ev in ev_list2:
                        excerpt = (ev or {}).get("normalized_text") or (ev or {}).get("text") or ""
                        if excerpt:
                            evidence_text = excerpt
                            break
                if evidence_text:
                    sub_blocks = _split_contract_blob(
                        clause_rec.normalized_text, clause_uid, clause_rec.page_number
                    )
                    needle = evidence_text[:120].lower()
                    for sub_uid, _sub_style, sub_text in sub_blocks:
                        if needle[:60] in sub_text.lower():
                            target_uid = sub_uid
                            break

            edits: dict = dict(run.contract_edits or {})
            edits[target_uid] = {
                "schema_version": 2,
                "text": replacement_text,
                "plain_text": _safe_rich_text_to_plain(replacement_text),
                "rich_text": [],
                "page": None,
                "rects": [],
                "anchor_text": None,
                "edited_at": datetime.utcnow().isoformat(),
            }
            run.contract_edits = edits
            flag_modified(run, "contract_edits")
            applied_text = replacement_text

    # Mark finding as accepted so it disappears from the review queue
    if finding:
        finding.accepted_at = datetime.utcnow()
    else:
        # admin_merge-only finding (no FindingRecord row) — persist accepted id on the run
        _run = (await db.execute(select(RunRecord).where(RunRecord.id == run_id))).scalar_one_or_none()
        if _run is not None:
            payload: dict = dict(_run.verdict_payload or {})
            accepted_ids: list = list(payload.get("accepted_admin_finding_ids") or [])
            if finding_id not in accepted_ids:
                accepted_ids.append(finding_id)
            payload["accepted_admin_finding_ids"] = accepted_ids
            _run.verdict_payload = payload
            flag_modified(_run, "verdict_payload")

    annotation_result = await db.execute(
        select(DocumentAnnotationRecord).where(
            DocumentAnnotationRecord.run_id == run_id,
            DocumentAnnotationRecord.finding_id == finding_id,
            DocumentAnnotationRecord.status == "open",
        )
    )
    for annotation in annotation_result.scalars().all():
        annotation.status = "accepted"

    # Audit trail
    if finding:
        record = HumanReviewRecord(
            id=str(uuid.uuid4()),
            run_id=run_id,
            finding_id=finding_id,
            reviewer_id=token["sub"],
            reviewer_role="human_reviewer",
            action="edit",
            original_finding_snapshot={"issue": finding.issue, "severity": finding.severity},
            edited_finding_snapshot={
                "accepted_recommendation": finding.recommendation,
                "applied_text": applied_text,
            },
            edit_justification="accepted",
            is_run_level=False,
        )
        db.add(record)
    await db.commit()
    return {"finding_id": finding_id, "accepted": True, "applied_text": applied_text}


@router.post("/runs/{run_id}/findings/{finding_id}/dismiss", status_code=status.HTTP_200_OK)
async def dismiss_finding(
    run_id: str,
    finding_id: str,
    db: DbSession,
    token: dict = Depends(require_auth),
) -> dict:
    finding = (await db.execute(
        select(FindingRecord).where(FindingRecord.id == finding_id, FindingRecord.run_id == run_id)
    )).scalar_one_or_none()
    if not finding:
        # admin_merge-only finding (no FindingRecord row) — persist dismissed id on the run.
        from app.backend.services.run_service import _load_full_findings
        if not any(f.finding_id == finding_id for f in await _load_full_findings(db, run_id)):
            raise HTTPException(status_code=404, detail="Finding not found")
        _run = (await db.execute(select(RunRecord).where(RunRecord.id == run_id))).scalar_one_or_none()
        if _run is not None:
            payload: dict = dict(_run.verdict_payload or {})
            dismissed_ids: list = list(payload.get("dismissed_admin_finding_ids") or [])
            if finding_id not in dismissed_ids:
                dismissed_ids.append(finding_id)
            payload["dismissed_admin_finding_ids"] = dismissed_ids
            _run.verdict_payload = payload
            flag_modified(_run, "verdict_payload")
            await db.commit()
        return {"finding_id": finding_id, "dismissed": True}

    finding.dismissed_at = datetime.utcnow()
    annotation_result = await db.execute(
        select(DocumentAnnotationRecord).where(
            DocumentAnnotationRecord.run_id == run_id,
            DocumentAnnotationRecord.finding_id == finding_id,
            DocumentAnnotationRecord.status == "open",
        )
    )
    for annotation in annotation_result.scalars().all():
        annotation.status = "dismissed"

    # Audit trail
    record = HumanReviewRecord(
        id=str(uuid.uuid4()),
        run_id=run_id,
        finding_id=finding_id,
        reviewer_id=token["sub"],
        reviewer_role="human_reviewer",
        action="reject",
        original_finding_snapshot={"issue": finding.issue, "severity": finding.severity},
        is_run_level=False,
    )
    db.add(record)
    await db.commit()
    return {"finding_id": finding_id, "dismissed": True}


# ---------------------------------------------------------------------------
# Document annotations — anchored comments and suggestions
# ---------------------------------------------------------------------------


def _annotation_to_out(a: DocumentAnnotationRecord) -> DocumentAnnotationOut:
    author_name = None
    if a.author:
        author_name = getattr(a.author, "display_name", None)
    return DocumentAnnotationOut(
        id=a.id,
        run_id=a.run_id,
        clause_uid=a.clause_uid,
        page_number=a.page_number,
        selected_text=a.selected_text,
        span_start=a.span_start,
        span_end=a.span_end,
        annotation_type=a.annotation_type,
        body=a.body,
        suggested_replacement=a.suggested_replacement,
        status=a.status,
        source=a.source,
        finding_id=a.finding_id,
        author_id=a.author_id,
        author_name=author_name,
        created_at=a.created_at,
        updated_at=a.updated_at,
    )


@router.get("/runs/{run_id}/annotations", response_model=list[DocumentAnnotationOut])
async def list_annotations(
    run_id: str,
    db: DbSession,
    clause_uid: str | None = None,
) -> list[DocumentAnnotationOut]:
    q = (
        select(DocumentAnnotationRecord)
        .options(selectinload(DocumentAnnotationRecord.author))
        .where(
            DocumentAnnotationRecord.run_id == run_id,
            DocumentAnnotationRecord.status.not_in(["accepted", "dismissed", "deleted"]),
        )
        .order_by(DocumentAnnotationRecord.created_at.asc())
    )
    if clause_uid:
        q = q.where(DocumentAnnotationRecord.clause_uid == clause_uid)
    rows = (await db.execute(q)).scalars().all()
    return [_annotation_to_out(a) for a in rows]


@router.post("/runs/{run_id}/annotations", response_model=DocumentAnnotationOut, status_code=status.HTTP_201_CREATED)
async def create_annotation(
    run_id: str,
    body: DocumentAnnotationCreate,
    db: DbSession,
    token: dict = Depends(require_auth),
) -> DocumentAnnotationOut:
    if body.annotation_type not in ("comment", "suggestion"):
        raise HTTPException(status_code=400, detail="annotation_type must be comment or suggestion")
    if body.annotation_type == "suggestion" and not (body.suggested_replacement or "").strip():
        raise HTTPException(status_code=400, detail="Suggestions require suggested_replacement")
    record = DocumentAnnotationRecord(
        id=str(uuid.uuid4()),
        run_id=run_id,
        clause_uid=body.clause_uid,
        annotation_type=body.annotation_type,
        body=body.body.strip(),
        suggested_replacement=body.suggested_replacement,
        selected_text=body.selected_text,
        span_start=body.span_start,
        span_end=body.span_end,
        page_number=body.page_number,
        source="human",
        author_id=token["sub"],
        status="open",
    )
    db.add(record)
    await db.flush()
    await db.refresh(record, ["author"])
    return _annotation_to_out(record)


@router.delete("/runs/{run_id}/annotations/{annotation_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_annotation(
    run_id: str,
    annotation_id: str,
    db: DbSession,
    token: dict = Depends(require_auth),
) -> None:
    record = (await db.execute(
        select(DocumentAnnotationRecord).where(
            DocumentAnnotationRecord.id == annotation_id,
            DocumentAnnotationRecord.run_id == run_id,
        )
    )).scalar_one_or_none()
    if not record or record.status == "deleted":
        raise HTTPException(status_code=404, detail="Annotation not found")
    if record.source == "human" and record.author_id != token["sub"]:
        raise HTTPException(status_code=403, detail="Cannot delete another user's annotation")
    record.status = "deleted"
    await db.commit()


# ---------------------------------------------------------------------------
# Document draft — computed from existing parsed data, no extra table needed
# ---------------------------------------------------------------------------

_HEADING1_RE = re.compile(r"^(ARTICLE|SECTION|SCHEDULE|EXHIBIT|ANNEX|PART|CHAPTER)\s+", re.I)
_HEADING2_NUM_RE = re.compile(r"^\d+\.\s+[A-Z]")
_HEADING3_NUM_RE = re.compile(r"^\d+\.\d+[\s.]+")
_LIST_RE = re.compile(r"^\([a-z]\)\s|^[a-z]\)\s+|\(i[ivx]*\)\s")


def _infer_block_style(text: str) -> str:
    t = text.strip()
    if not t:
        return "body"
    if t.isupper() and 3 < len(t) <= 70:
        return "heading1"
    if _HEADING1_RE.match(t) and len(t) <= 150:
        return "heading1"
    # Tighten heading2/3 to ≤75 chars so "7. Governing Law This Agreement shall…"
    # (heading title + body text inlined) falls through to body style instead of bold.
    if _HEADING2_NUM_RE.match(t) and len(t) <= 75 and not t.endswith(","):
        return "heading2"
    if _HEADING3_NUM_RE.match(t) and len(t) <= 75:
        return "heading3"
    if _LIST_RE.match(t):
        return "list_item"
    return "body"



# ---------------------------------------------------------------------------
# Contract blob splitter — breaks a single giant clause into styled sub-blocks
# ---------------------------------------------------------------------------

_HEADING2_BODY_SPLIT = re.compile(
    r'^(\d+\.(?:\d+\.)?\s+[A-Z][A-Za-z ,\-–—]{1,80}?)\s+(?=(?:The|This|For|A\b|An\b|Any\b|Each\b|Upon|All\b|No\b|In\b|If\b|When\b|Where\b|Except|Subject|Notwithstanding|It\b|At\b|By\b|Such\b|Under\b|Pursuant|With\b|Without|Recipient|Neither)\b)'
)


def _split_contract_blob(text: str, base_uid: str, page_number: int) -> list[tuple[str, str, str]]:
    """Split a large contract blob into (sub_uid, style, sub_text) tuples."""
    results: list[tuple[str, str, str]] = []
    counter = [0]

    def next_uid(first: bool = False) -> str:
        if first:
            return base_uid
        counter[0] += 1
        return f"{base_uid}__p{counter[0]}"

    remaining = text.strip()
    if not remaining:
        return [(base_uid, "body", text)]

    # 1. Extract leading ALL-CAPS title
    caps_m = re.match(r'^([A-Z][A-Z\s\-–—\(\)]{3,80}?)(?=\s+[A-Z][a-z]|\s+[a-z0-9])', remaining)
    if caps_m:
        title = caps_m.group(1).strip()
        if not re.search(r'[a-z]', title):  # truly all caps
            uid = next_uid(first=len(results) == 0)
            results.append((uid, "heading1", title))
            remaining = remaining[caps_m.end():].strip()

    # 2. Extract subtitle (short mixed-case before body)
    if remaining:
        sub_m = re.match(
            r'^([A-Z][A-Za-z0-9\-\s]{3,55}?)(?=\s+(?:This|The|A\b|An\b|For\b|In\b|These\b|Each\b|It\b))',
            remaining
        )
        if sub_m and '.' not in sub_m.group(1) and not re.match(r'^\d+\.', sub_m.group(1)) and len(sub_m.group(1)) <= 60:
            subtitle = sub_m.group(1).strip()
            uid = next_uid(first=len(results) == 0)
            results.append((uid, "subtitle", subtitle))
            remaining = remaining[sub_m.end():].strip()

    # 3. Split on "IN WITNESS WHEREOF" boundary
    witness_marker = "IN WITNESS WHEREOF"
    witness_idx = remaining.find(witness_marker)
    if witness_idx >= 0:
        main_body = remaining[:witness_idx].strip()
        witness_block = remaining[witness_idx:].strip()
    else:
        main_body = remaining
        witness_block = None

    # 4. Split main body on numbered section patterns
    if main_body:
        parts = re.split(r'(?<=[."\'\)])\s+(?=\d+\.\s+[A-Z])', main_body)
        for part in parts:
            part = part.strip()
            if not part:
                continue
            # 5. For each part try to split heading from body
            hm = _HEADING2_BODY_SPLIT.match(part)
            if hm:
                heading_text = hm.group(1).strip()
                body_text = part[hm.end():].strip()
                uid = next_uid(first=len(results) == 0)
                results.append((uid, "heading2", heading_text))
                if body_text:
                    uid = next_uid(first=len(results) == 0)
                    results.append((uid, "body", body_text))
            else:
                uid = next_uid(first=len(results) == 0)
                results.append((uid, "body", part))

    # Add witness block as body
    if witness_block:
        uid = next_uid(first=len(results) == 0)
        results.append((uid, "body", witness_block))

    # Fallback: if nothing was split, return single body block
    if not results:
        return [(base_uid, "body", text)]

    return results


def _compute_comment_anchors(original_text: str, anns: list) -> list[CommentAnchorOut]:
    result = []
    lower_orig = original_text.lower()
    for ann in anns:
        if not ann.selected_text:
            continue
        span_start = getattr(ann, "span_start", None)
        span_end = getattr(ann, "span_end", None)
        # Validate stored positions against selected_text before trusting them
        if span_start is not None and span_end is not None and span_end <= len(original_text):
            sliced = original_text[span_start:span_end]
            norm_sliced = " ".join(sliced.lower().split())
            norm_sel_chk = " ".join(ann.selected_text.lower().split())
            if norm_sliced == norm_sel_chk:
                result.append(CommentAnchorOut(annotation_id=ann.id, from_pos=span_start, to_pos=span_end))
                continue
        # Fallback: case-insensitive substring search in original text
        sel_stripped = ann.selected_text.strip()
        lower_sel = sel_stripped.lower()
        idx = lower_orig.find(lower_sel)
        if idx >= 0:
            result.append(CommentAnchorOut(annotation_id=ann.id, from_pos=idx, to_pos=idx + len(sel_stripped)))
            continue
        # Last resort: normalized text search (positions in normalized space)
        norm_orig = " ".join(original_text.lower().split())
        norm_sel = " ".join(ann.selected_text.lower().split())
        idx = norm_orig.find(norm_sel)
        if idx >= 0:
            result.append(CommentAnchorOut(annotation_id=ann.id, from_pos=idx, to_pos=idx + len(norm_sel)))
    return result


@router.get("/runs/{run_id}/document-draft", response_model=DocumentDraftOut)
async def get_document_draft(run_id: str, db: DbSession) -> DocumentDraftOut:
    """Return structured document blocks for the Google-Docs-style editor."""
    run = (await db.execute(select(RunRecord).where(RunRecord.id == run_id))).scalar_one_or_none()
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")

    clauses = (await db.execute(
        select(ParsedClauseRecord)
        .where(ParsedClauseRecord.run_id == run_id)
        .order_by(ParsedClauseRecord.order_index.asc().nullslast(), ParsedClauseRecord.page_number.asc(), ParsedClauseRecord.bbox_y0.asc())
    )).scalars().all()

    annotations = (await db.execute(
        select(DocumentAnnotationRecord).where(
            DocumentAnnotationRecord.run_id == run_id,
            DocumentAnnotationRecord.status.not_in(["accepted", "dismissed", "deleted"]),
        )
    )).scalars().all()
    ann_by_clause: dict[str, list] = {}
    for ann in annotations:
        ann_by_clause.setdefault(ann.clause_uid, []).append(ann)

    # Pending findings: use _load_full_findings so admin_merge + DB findings both surface.
    # This ensures findings from structured stage output (no FindingRecord row) still render
    # as pending suggestions when they carry a recommended_change.
    from app.backend.services.run_service import _load_full_findings as _load_findings_for_draft
    all_pending = await _load_findings_for_draft(db, run_id)
    finding_by_clause: dict[str, object] = {}
    for f in all_pending:
        if f.clause_uid and f.clause_uid not in finding_by_clause and (f.recommended_change or "").strip():
            finding_by_clause[f.clause_uid] = f

    edits: dict = run.contract_edits or {}
    revision = len(edits)

    # ------------------------------------------------------------------
    # DOCX-based path: use pdf2docx cache when available so that bold,
    # italic, and paragraph structure from the original PDF are preserved.
    # ------------------------------------------------------------------
    docx_cache: "Path | None" = None
    if run.storage_path:
        cand = Path(run.storage_path).with_suffix(".original.docx")
        if cand.exists():
            docx_cache = cand

    if docx_cache is not None:
        try:
            docx_paragraphs = _extract_docx_blocks(str(docx_cache))
        except Exception:
            docx_paragraphs = []

        if docx_paragraphs:
            logger.info("document-draft: run=%s docx_cache_path_exists=%s blocks_built_via=docx", run_id, docx_cache.exists() if docx_cache else False)
            # Build a lookup: norm_prefix_60 → clause record (first match wins)
            clause_by_prefix: dict[str, object] = {}
            for clause in clauses:
                key = _norm_text(clause.normalized_text)[:60]
                if key not in clause_by_prefix:
                    clause_by_prefix[key] = clause

            blocks: list[DraftBlockOut] = []
            last_clause = clauses[0] if clauses else None
            for para_idx, para in enumerate(docx_paragraphs):
                para_text: str = para["text"]
                para_style: str = para["style"]
                para_marks: list[dict] = para["marks"]

                # Match paragraph to a clause by normalized-text prefix
                para_norm = _norm_text(para_text)[:60]
                matched_clause = clause_by_prefix.get(para_norm)
                if matched_clause is not None:
                    last_clause = matched_clause
                current_clause = last_clause

                if current_clause is None:
                    continue

                clause_uid: str = current_clause.clause_uid
                page_number: int = current_clause.page_number

                # Apply accepted edit (plain text overrides marks)
                raw_edit = edits.get(clause_uid)
                if raw_edit and isinstance(raw_edit, dict):
                    display_text = raw_edit.get("plain_text") or _safe_rich_text_to_plain(raw_edit.get("text", "")) or para_text
                    display_marks: list[dict] = []
                else:
                    display_text = para_text
                    display_marks = para_marks

                # Pending suggestion — only on body/list_item blocks
                pending = finding_by_clause.get(clause_uid)
                pending_out: Optional[PendingSuggestionOut] = None
                if pending and para_style in ("body", "list_item"):
                    # Only emit on the FIRST paragraph whose normalized text contains
                    # the finding's evidence text (evidence-match like _find_pending_sub_uid)
                    evidence_text = ""
                    for ev in (getattr(pending, "contract_evidence", None) or []):
                        t = (getattr(ev, "text", "") or "").strip()
                        if t:
                            evidence_text = t
                            break
                    emit_suggestion = False
                    if evidence_text:
                        needle = evidence_text[:60].lower()
                        emit_suggestion = needle in para_text.lower()
                    else:
                        # Fallback: emit on the first body block for this clause
                        emit_suggestion = (matched_clause is not None)
                    if emit_suggestion:
                        pending_out = PendingSuggestionOut(
                            finding_id=pending.finding_id,
                            severity=pending.severity or "medium",
                            description=pending.description or "",
                            replacement_text=(pending.recommended_change or "").strip(),
                        )

                # Comment anchors
                clause_anns = ann_by_clause.get(clause_uid, [])
                comment_anchors = _compute_comment_anchors(para_text, clause_anns)

                block_id = f"block_{clause_uid}_{para_idx}"
                blocks.append(DraftBlockOut(
                    block_id=block_id,
                    clause_uid=clause_uid,
                    page_number=page_number,
                    style=para_style,
                    text=display_text,
                    original_text=para_text,
                    marks=display_marks,
                    pending_suggestion=pending_out,
                    comment_anchors=comment_anchors,
                ))

            return DocumentDraftOut(blocks=blocks, revision=revision)

    # ------------------------------------------------------------------
    # Fallback: clause-iteration path (no DOCX cache available)
    # ------------------------------------------------------------------
    logger.info("document-draft: run=%s docx_cache_path_exists=%s blocks_built_via=clauses", run_id, docx_cache.exists() if docx_cache else False)
    blocks: list[DraftBlockOut] = []
    for clause in clauses:
        raw_edit = edits.get(clause.clause_uid)
        if raw_edit and isinstance(raw_edit, dict):
            current_text = raw_edit.get("plain_text") or _safe_rich_text_to_plain(raw_edit.get("text", ""))
        else:
            current_text = clause.normalized_text
        original_text = clause.normalized_text

        # For large blobs (>300 chars), split into styled sub-blocks
        if len(original_text) > 300:
            sub_blocks = _split_contract_blob(original_text, clause.clause_uid, clause.page_number)
            clause_anns = ann_by_clause.get(clause.clause_uid, [])
            pending = finding_by_clause.get(clause.clause_uid)

            # Pre-compute which sub-block should show the pending suggestion
            pending_target_uid = _find_pending_sub_uid(sub_blocks, pending) if pending else None

            for sub_uid, sub_style, sub_text in sub_blocks:
                # Only use an edit if it's stored under the sub-block's own UID (i.e. __pN suffix).
                # Never apply the parent clause's legacy full-blob edit to any sub-block.
                is_sub_uid = sub_uid != clause.clause_uid
                sub_edit = edits.get(sub_uid) if is_sub_uid else None
                if sub_edit and isinstance(sub_edit, dict):
                    display_text = sub_edit.get("plain_text") or _safe_rich_text_to_plain(sub_edit.get("text", "")) or sub_text
                else:
                    display_text = sub_text

                # Filter comment anchors to those within this sub-block's offset range.
                # Also include annotations stored under the sub-block's own UID (span coords
                # are already relative to sub_text so need no offset adjustment).
                sub_uid_anns = ann_by_clause.get(sub_uid, [])
                sub_offset = original_text.find(sub_text)
                if sub_offset >= 0:
                    sub_end = sub_offset + len(sub_text)
                    sub_anns = [
                        ann for ann in clause_anns
                        if ann.span_start is not None and ann.span_end is not None
                        and ann.span_start >= sub_offset and ann.span_end <= sub_end
                    ]
                    adjusted_anns = []
                    for ann in sub_anns:
                        class _AdjustedAnn:
                            pass
                        adj = _AdjustedAnn()
                        adj.id = ann.id
                        adj.selected_text = ann.selected_text
                        adj.span_start = ann.span_start - sub_offset
                        adj.span_end = ann.span_end - sub_offset
                        adjusted_anns.append(adj)
                    # Annotations saved directly under sub_uid need no offset adjustment
                    adjusted_anns.extend(sub_uid_anns)
                    comment_anchors = _compute_comment_anchors(sub_text, adjusted_anns)
                else:
                    comment_anchors = _compute_comment_anchors(sub_text, sub_uid_anns)

                # Show pending suggestion only on the evidence-matched sub-block
                pending_out: Optional[PendingSuggestionOut] = None
                if pending and sub_uid == pending_target_uid:
                    pending_out = PendingSuggestionOut(
                        finding_id=pending.finding_id,
                        severity=pending.severity or "medium",
                        description=pending.description or "",
                        replacement_text=(pending.recommended_change or "").strip(),
                    )

                blocks.append(DraftBlockOut(
                    block_id=f"block_{sub_uid}",
                    clause_uid=sub_uid,
                    page_number=clause.page_number,
                    style=sub_style,
                    text=display_text,
                    original_text=sub_text,
                    marks=[],
                    pending_suggestion=pending_out,
                    comment_anchors=comment_anchors,
                ))
        else:
            # Short clause: existing single-block logic
            style = _infer_block_style(original_text)

            pending = finding_by_clause.get(clause.clause_uid)
            pending_out = None
            if pending and style not in ("heading1", "heading2", "heading3"):
                pending_out = PendingSuggestionOut(
                    finding_id=pending.finding_id,
                    severity=pending.severity or "medium",
                    description=pending.description or "",
                    replacement_text=(pending.recommended_change or "").strip(),
                )

            clause_anns = ann_by_clause.get(clause.clause_uid, [])
            comment_anchors = _compute_comment_anchors(original_text, clause_anns)

            blocks.append(DraftBlockOut(
                block_id=f"block_{clause.clause_uid}",
                clause_uid=clause.clause_uid,
                page_number=clause.page_number,
                style=style,
                text=current_text,
                original_text=original_text,
                marks=[],
                pending_suggestion=pending_out,
                comment_anchors=comment_anchors,
            ))

    return DocumentDraftOut(blocks=blocks, revision=revision)


# ---------------------------------------------------------------------------
# DOCX export — export clean document draft as a Word file
# ---------------------------------------------------------------------------


@router.get("/runs/{run_id}/export-docx")
async def export_docx(run_id: str, db: DbSession, original: bool = False) -> StreamingResponse:
    """Export document as DOCX. ?original=true returns pdf2docx-converted layout-preserved original."""
    run = (await db.execute(select(RunRecord).where(RunRecord.id == run_id))).scalar_one_or_none()
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")

    base = (run.original_filename or "contract").rsplit(".", 1)[0]

    if original:
        path = Path(run.storage_path)
        if path.exists():
            try:
                from pdf2docx import Converter as _Pdf2DocxConverter
                docx_cache = path.with_suffix(".original.docx")
                logger.info("export-docx: run=%s original=%s cache_exists=%s", run_id, original, docx_cache.exists())
                if not docx_cache.exists():
                    cv = _Pdf2DocxConverter(str(path))
                    cv.convert(str(docx_cache), start=0, end=None)
                    cv.close()
                buf = io.BytesIO(docx_cache.read_bytes())
                buf.seek(0)
                return StreamingResponse(
                    buf,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f'attachment; filename="{base}.docx"'},
                )
            except Exception:
                pass  # fall through to parsed-clause approach

    clauses = (await db.execute(
        select(ParsedClauseRecord)
        .where(ParsedClauseRecord.run_id == run_id)
        .order_by(ParsedClauseRecord.order_index.asc().nullslast(), ParsedClauseRecord.page_number.asc(), ParsedClauseRecord.bbox_y0.asc())
    )).scalars().all()

    edits: dict = {} if original else (run.contract_edits or {})
    doc = _build_modified_docx(run, list(clauses), edits)
    suffix = "" if original else "-edited"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{base}{suffix}.docx"'},
    )
