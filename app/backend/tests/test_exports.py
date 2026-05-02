"""
Tests for clean DOCX and PDF export: accepted edits appear as plain text,
deleted text is absent, no tracked-change XML, no redline markers.
"""
from __future__ import annotations

import io
import types
import zipfile

import pytest


# ---------------------------------------------------------------------------
# Minimal stubs
# ---------------------------------------------------------------------------


def _make_clause(uid: str, text: str, page: int = 1):
    c = types.SimpleNamespace()
    c.clause_uid = uid
    c.normalized_text = text
    c.page_number = page
    c.bbox_y0 = 0.0
    return c


def _make_edit(plain_text: str) -> dict:
    return {
        "schema_version": 2,
        "text": plain_text,
        "plain_text": plain_text,
        "rich_text": [],
        "page": 1,
        "rects": [],
        "anchor_text": None,
        "edited_at": "2026-05-02T00:00:00",
    }


def _import_helpers():
    from app.backend.routes.editor import (
        _build_clean_pdf,
        _build_document_blocks,
        _find_pending_sub_uid,
        _split_contract_blob,
    )
    return _build_document_blocks, _build_clean_pdf, _find_pending_sub_uid, _split_contract_blob


# ---------------------------------------------------------------------------
# _build_document_blocks
# ---------------------------------------------------------------------------


class TestBuildDocumentBlocks:
    def test_accepted_edit_replaces_original(self):
        _build_document_blocks, *_ = _import_helpers()
        clause = _make_clause("c1", "The original clause text.")
        edits = {"c1": _make_edit("The replaced clause text.")}
        blocks = _build_document_blocks([clause], edits)
        texts = [t for _, t in blocks]
        assert any("replaced" in t for t in texts), "accepted edit text must appear"
        assert not any("original" in t for t in texts), "original text must be absent when edited"

    def test_deleted_text_absent_when_replaced(self):
        _build_document_blocks, *_ = _import_helpers()
        clause = _make_clause("c1", "Seller warrants title to the property.")
        edits = {"c1": _make_edit("Seller makes no warranty as to title.")}
        blocks = _build_document_blocks([clause], edits)
        texts = [t for _, t in blocks]
        assert not any("warrants title" in t for t in texts)

    def test_original_mode_ignores_edits(self):
        _build_document_blocks, *_ = _import_helpers()
        clause = _make_clause("c1", "The original clause text.")
        edits = {"c1": _make_edit("The replaced clause text.")}
        blocks = _build_document_blocks([clause], edits, original=True)
        texts = [t for _, t in blocks]
        assert any("original" in t for t in texts)
        assert not any("replaced" in t for t in texts)

    def test_large_blob_sub_block_edit_applied(self):
        _build_document_blocks, _, __, _split = _import_helpers()
        long_text = "INDEMNIFICATION " + ("The indemnifying party shall defend. " * 20)
        clause = _make_clause("c1", long_text)
        sub_blocks = _split(long_text, "c1", 1)
        body_uid = next((uid for uid, style, _ in sub_blocks if style in ("body", "list_item")), None)
        assert body_uid is not None
        edits = {body_uid: _make_edit("Indemnification obligations are hereby limited.")}
        blocks = _build_document_blocks([clause], edits)
        texts = [t for _, t in blocks]
        assert any("limited" in t for t in texts)


# ---------------------------------------------------------------------------
# _build_clean_pdf
# ---------------------------------------------------------------------------


class TestBuildCleanPdf:
    def test_returns_pdf_bytes(self):
        _, _build_clean_pdf, *_ = _import_helpers()
        blocks = [("heading1", "SERVICE AGREEMENT"), ("body", "This agreement is entered into.")]
        buf = _build_clean_pdf(blocks)
        data = buf.read()
        assert data[:4] == b"%PDF", "output must be valid PDF"

    def test_no_replacement_label_in_pdf(self):
        _, _build_clean_pdf, *_ = _import_helpers()
        blocks = [("body", "Clean final contract clause.")]
        buf = _build_clean_pdf(blocks)
        data = buf.read()
        assert b"REPLACEMENT" not in data, "clean PDF must not contain REPLACEMENT label"

    def test_no_redline_color_commands(self):
        """Old overlay used '0.8 0.1 0.1 RG' for red strikethrough lines."""
        _, _build_clean_pdf, *_ = _import_helpers()
        blocks = [("body", "Final clause without redlines.")]
        buf = _build_clean_pdf(blocks)
        data = buf.read()
        assert b"0.8 0.1 0.1 RG" not in data

    def test_no_green_background_commands(self):
        """Old overlay used '0.94 0.99 0.95 rg' for green replacement box fill."""
        _, _build_clean_pdf, *_ = _import_helpers()
        blocks = [("body", "Final clause without green replacement boxes.")]
        buf = _build_clean_pdf(blocks)
        data = buf.read()
        assert b"0.94 0.99 0.95 rg" not in data

    def test_empty_blocks_produces_valid_pdf(self):
        _, _build_clean_pdf, *_ = _import_helpers()
        buf = _build_clean_pdf([])
        data = buf.read()
        assert data[:4] == b"%PDF"


# ---------------------------------------------------------------------------
# DOCX: no tracked-change XML, accepted text present, deleted text absent
# ---------------------------------------------------------------------------


def _build_test_docx(clauses, edits) -> io.BytesIO:
    """Build a DOCX the same way the export-docx endpoint does, for testing."""
    import re
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn as _qn
    from app.backend.routes.editor import (
        _build_document_blocks,
        _infer_block_style,
        _safe_rich_text_to_plain,
    )

    doc = DocxDocument()
    try:
        settings_el = doc.settings.element
        for tag in ("w:trackChanges", "w:revisionView"):
            for el in settings_el.findall(_qn(tag)):
                settings_el.remove(el)
    except Exception:
        pass
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    blocks = _build_document_blocks(clauses, edits)
    for style_key, text in blocks:
        p = doc.add_paragraph()
        run_obj = p.add_run(text)
        run_obj.font.name = "Times New Roman"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _docx_xml(buf: io.BytesIO) -> bytes:
    with zipfile.ZipFile(buf) as z:
        return z.read("word/document.xml")


class TestDocxNoTrackedChanges:
    def test_no_w_ins_elements(self):
        clause = _make_clause("c1", "This is the original clause.")
        edits = {"c1": _make_edit("This is the accepted replacement.")}
        xml = _docx_xml(_build_test_docx([clause], edits))
        assert b"w:ins" not in xml, "DOCX must not contain w:ins tracked-insertion elements"

    def test_no_w_del_elements(self):
        clause = _make_clause("c1", "This is the original clause.")
        edits = {"c1": _make_edit("This is the accepted replacement.")}
        xml = _docx_xml(_build_test_docx([clause], edits))
        assert b"w:del" not in xml, "DOCX must not contain w:del tracked-deletion elements"

    def test_accepted_text_present_in_docx_xml(self):
        clause = _make_clause("c1", "The original text.")
        edits = {"c1": _make_edit("The accepted replacement.")}
        xml = _docx_xml(_build_test_docx([clause], edits))
        assert b"accepted replacement" in xml

    def test_original_text_absent_in_modified_docx(self):
        clause = _make_clause("c1", "The original text.")
        edits = {"c1": _make_edit("The accepted replacement.")}
        xml = _docx_xml(_build_test_docx([clause], edits))
        assert b"original text" not in xml


# ---------------------------------------------------------------------------
# _find_pending_sub_uid
# ---------------------------------------------------------------------------


class TestFindPendingSubUid:
    def test_matches_by_evidence_text(self):
        _, __, _find_pending_sub_uid, _split = _import_helpers()
        long_text = "LIMITATION OF LIABILITY The total liability shall not exceed the contract value. " * 5
        sub_blocks = _split(long_text, "c1", 1)
        finding = types.SimpleNamespace(
            contract_evidence=[types.SimpleNamespace(text="total liability shall not exceed")]
        )
        uid = _find_pending_sub_uid(sub_blocks, finding)
        assert uid is not None
        matched_text = next(t for u, _, t in sub_blocks if u == uid)
        assert "total liability" in matched_text

    def test_fallback_to_first_body_block(self):
        _, __, _find_pending_sub_uid, _split = _import_helpers()
        long_text = "INDEMNIFICATION " + ("indemnifying party shall defend. " * 10)
        sub_blocks = _split(long_text, "c1", 1)
        finding = types.SimpleNamespace(contract_evidence=[])
        uid = _find_pending_sub_uid(sub_blocks, finding)
        assert uid is not None
        style = next(s for u, s, _ in sub_blocks if u == uid)
        assert style in ("body", "list_item")

    def test_none_when_no_blocks(self):
        _, __, _find_pending_sub_uid, _ = _import_helpers()
        finding = types.SimpleNamespace(contract_evidence=[])
        uid = _find_pending_sub_uid([], finding)
        assert uid is None
